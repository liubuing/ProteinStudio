#!/usr/bin/env python
"""Generate Word user manual for the Protein Sequence Design Platform."""
import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DESKTOP = os.path.join(os.environ['USERPROFILE'], 'Desktop')
OUT_PATH = os.path.join(DESKTOP, '蛋白质序列设计平台_使用说明.docx')

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading_style.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

# ═══════════════ Helper functions ═══════════════
def add_code_block(doc, text):
    """Add a code-formatted paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table_simple(doc, headers, rows):
    """Add a styled table."""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

def add_note(doc, text, note_type="info"):
    """Add a highlighted note."""
    p = doc.add_paragraph()
    prefixes = {"info": "ℹ️ ", "warn": "⚠️ ", "tip": "💡 "}
    prefix = prefixes.get(note_type, "")
    run = p.add_run(prefix + text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    if note_type == "warn":
        run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
    elif note_type == "tip":
        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    return p

# ═══════════════ Title Page ═══════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('蛋白质序列设计平台')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('使用说明手册')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('基于 BFN · ProteinMPNN · ESM-IF · AlphaFold2\nGradio Web 界面 · 2026年5月')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════ Table of Contents ═══════════════
doc.add_heading('目录', level=1)
toc_items = [
    '1. 平台概述',
    '2. 快速启动',
    '3. FASTA → PDB 转换（α-螺旋模板）',
    '4. 单步设计（Tab 1）',
    '5. 批量处理（Tab 2）',
    '6. 结构预测（Tab 3 — AlphaFold2）',
    '7. 系统设置（Tab 4）',
    '8. 靶点设计（Tab 5）',
    '9. 迭代精修（Tab 6）',
    '10. 置信度指标说明',
    '11. 级联过滤机制',
    '12. 配置文件参考',
    '13. 命令行工具',
    '14. 8种测试蛋白质结果',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

doc.add_page_break()

# ═══════════════ 1. Platform Overview ═══════════════
doc.add_heading('1. 平台概述', level=1)
doc.add_paragraph(
    '本平台集成了三种蛋白质序列设计工具和一种结构预测工具，通过 Gradio Web 界面提供'
    '统一的操作体验。核心功能包括：从序列生成α-螺旋模板、单步/批量蛋白质设计、'
    '靶点导向的抗体设计、AF2驱动的迭代精修。'
)

doc.add_heading('1.1 设计工具', level=2)
add_table_simple(doc,
    ['工具', '全称', '适用场景', '输入要求'],
    [
        ['BFN (抗体CDR)', 'Bayesian Flow Network', '抗体CDR环区设计', '含H/L链的抗体PDB'],
        ['BFN (通用蛋白)', 'Bayesian Flow Network', '通用蛋白质序列设计', '任意蛋白质PDB'],
        ['ProteinMPNN', 'Message-Passing NN', '蛋白质骨架序列设计', '含N/CA/C/O骨架的PDB'],
        ['ESM-IF', 'ESM Inverse Folding', '结构→序列逆向设计', '含完整主链的PDB'],
    ]
)

doc.add_heading('1.2 结构预测', level=2)
doc.add_paragraph(
    'AlphaFold2 (ColabFold) — 给定氨基酸序列，预测三维结构。支持 pTM/ipTM/pLDDT/PAE '
    '置信度输出，用于下游验证和迭代精修。'
)

doc.add_heading('1.3 置信度指标', level=2)
add_table_simple(doc,
    ['指标', '来源', '范围', '含义'],
    [
        ['pLDDT', 'BFN / AF2', '0-1 (BFN) / 0-100 (AF2)', '逐残基预测置信度'],
        ['pTM', 'AF2', '0-1', '全局折叠置信度'],
        ['ipTM', 'AF2', '0-1', '界面/多链置信度 (multimer)'],
        ['PAE', 'BFN / AF2', '0-∞ Å', '预测对齐误差'],
        ['PPL', 'BFN', '>1', '序列困惑度 (越低越好)'],
    ]
)

doc.add_page_break()

# ═══════════════ 2. Quick Start ═══════════════
doc.add_heading('2. 快速启动', level=1)

doc.add_heading('2.1 启动服务', level=2)
doc.add_paragraph('在项目根目录下执行：')
add_code_block(doc, 'python app.py')
doc.add_paragraph('启动后浏览器自动打开 http://127.0.0.1:7860。仅本机访问，外网无法连接。')

doc.add_heading('2.2 管理命令', level=2)
add_code_block(doc, 'python manage.py start          # 启动服务')
add_code_block(doc, 'python manage.py stop           # 停止服务')
add_code_block(doc, 'python manage.py restart        # 重启服务')
add_code_block(doc, 'python manage.py status         # 查看状态')

doc.add_heading('2.3 最简工作流', level=2)
doc.add_paragraph(
    '① 在「FASTA → PDB 转换」输入序列 → 生成α-螺旋模板 PDB\n'
    '② 上传 PDB 到「单步设计」→ 选工具 (BFN 通用蛋白) → 点"开始设计"\n'
    '③ 查看设计结果 → 下载 FASTA 文件 → 可选发送到 AF2 验证'
)

doc.add_page_break()

# ═══════════════ 3. FASTA → PDB ═══════════════
doc.add_heading('3. FASTA → PDB 转换（α-螺旋模板）', level=1)

doc.add_paragraph(
    '输入氨基酸序列，自动生成标准α-螺旋骨架结构的 PDB 文件。生成的 PDB 包含完整的'
    '主链原子 (N, CA, C, O) 和 CB 侧链起点，兼容 BFN、ProteinMPNN、ESM-IF 三款设计工具。'
)

doc.add_heading('3.1 几何参数', level=2)
add_table_simple(doc,
    ['参数', '值', '说明'],
    [
        ['残基/圈', '3.6', 'α-螺旋标准螺距'],
        ['螺距 (Å)', '1.5', '每残基沿螺旋轴的上升距离'],
        ['半径 (Å)', '2.3', 'CA 原子到螺旋轴的距离'],
        ['CA-CA 距离 (Å)', '≈ 3.83', '相邻残基 CA 间距'],
        ['N-CA 键长 (Å)', '1.47', '标准蛋白质几何'],
        ['CA-C 键长 (Å)', '1.53', '标准蛋白质几何'],
        ['C-O 键长 (Å)', '1.23', '标准蛋白质几何'],
        ['N-CA-C 键角', '≈ 114.7°', '接近目标值 111°'],
        ['CA-C-O 键角', '≈ 120°', '标准几何'],
    ]
)

doc.add_heading('3.2 使用方式', level=2)
doc.add_paragraph(
    '在「单步设计」或「靶点设计」Tab 中，展开「FASTA → PDB 转换」折叠面板：\n'
    '1. 在文本框粘贴序列 (或上传 .fasta 文件)\n'
    '2. 点击「生成 α-螺旋模板 PDB」\n'
    '3. 生成的 PDB 自动填入上传区域，可立即用于设计'
)
add_note(doc, '长序列 (>50 aa) 将自动截取中间段作为设计区域。全长序列折叠需用 AF2 预测。', 'info')

doc.add_page_break()

# ═══════════════ 4. Single Design ═══════════════
doc.add_heading('4. 单步设计（Tab 1 — 🎯 单步设计）', level=1)

doc.add_paragraph('核心工作区，按三步流程完成蛋白质序列设计。')

doc.add_heading('4.1 步骤一：上传 PDB', level=2)
doc.add_paragraph(
    '支持直接上传 .pdb 文件，或通过 FASTA → PDB 转换生成。上传后自动检测链ID和残基范围。'
)

doc.add_heading('4.2 步骤二：选择工具与参数', level=2)

doc.add_heading('BFN 通用蛋白设计', level=3)
add_table_simple(doc,
    ['参数', '类型', '默认值', '说明'],
    [
        ['设计区域', '文本', 'A:10-25', '格式: 链:起始-结束，如 A:10-25'],
        ['序列数', '滑块 (1-20)', '3', '每轮生成的序列数量'],
        ['随机采样', '复选框', 'False', '是否随机采样 (vs 确定性)'],
        ['评估模式', '复选框', 'True', '使用训练时的评估模式'],
    ]
)

doc.add_heading('BFN 抗体 CDR 设计', level=3)
add_table_simple(doc,
    ['参数', '类型', '默认值', '说明'],
    [
        ['重链ID', '文本', 'H', '抗体重链标识符'],
        ['轻链ID', '文本', 'L', '抗体轻链标识符'],
        ['序列数', '滑块 (1-20)', '3', '每轮生成的序列数量'],
        ['随机采样', '复选框', 'False', '是否随机采样'],
        ['评估模式', '复选框', 'True', '使用训练时的评估模式'],
    ]
)

doc.add_heading('ProteinMPNN', level=3)
add_table_simple(doc,
    ['参数', '类型', '默认值', '说明'],
    [
        ['设计链', '文本', 'A', '需要设计的目标链ID'],
        ['温度', '文本', '0.1', '采样温度 (0.1-1.0)'],
        ['序列数', '滑块 (1-20)', '3', '生成的序列数量'],
        ['随机种子', '数字', '42', '随机数种子'],
        ['排除AA', '文本', '', '设计时排除的氨基酸, 如 "CX"'],
    ]
)

doc.add_heading('ESM-IF', level=3)
add_table_simple(doc,
    ['参数', '类型', '默认值', '说明'],
    [
        ['目标链', '文本', 'A', '需要设计的目标链ID'],
        ['温度', '滑块 (0.05-1.0)', '0.1', '采样温度'],
        ['序列数', '滑块 (1-10)', '3', '生成的序列数量'],
    ]
)

doc.add_heading('4.3 步骤三：结果', level=2)
doc.add_paragraph(
    '设计完成后显示：\n'
    '• 每条序列的氨基酸组成、PPL (困惑度)、pLDDT、ipTM、PAE\n'
    '• 级联过滤排名 (综合评分降序)\n'
    '• 下载 FASTA 文件按钮\n'
    '• 发送到 AlphaFold2 预测按钮 (一键切换至 Tab 3)'
)

add_note(doc, 'BFN 置信度头 (pLDDT/ipTM) 在通用蛋白上可能接近 0，这是正常现象。'
          '模型使用抗体数据进行训练，通用蛋白质的置信度输出未经过校准。请用 AF2 进行可靠验证。', 'warn')

doc.add_page_break()

# ═══════════════ 5. Batch ═══════════════
doc.add_heading('5. 批量处理（Tab 2 — 📦 批量处理）', level=1)

doc.add_paragraph(
    '对多个 PDB 文件批量运行同一设计工具，适合高通量筛选场景。'
)

doc.add_heading('5.1 参数说明', level=2)
add_table_simple(doc,
    ['参数', '说明'],
    [
        ['PDB 文件', '支持多文件上传 (.pdb)'],
        ['设计工具', 'BFN (通用蛋白) / ProteinMPNN / ESM-IF'],
        ['设计区域', '对 BFN 有效，格式同单步设计'],
        ['目标链', '对 MPNN/ESM-IF 有效'],
        ['温度', '采样温度参数'],
        ['每文件序列数', '每个 PDB 生成的序列数 (1-10)'],
    ]
)

doc.add_heading('5.2 输出', level=2)
doc.add_paragraph(
    '• 每个文件的独立设计结果\n'
    '• 汇总 JSON 文件 (含所有结果的结构化数据)\n'
    '• 控制台实时进度显示'
)

doc.add_page_break()

# ═══════════════ 6. Structure Prediction ═══════════════
doc.add_heading('6. 结构预测（Tab 3 — 🔮 结构预测）', level=1)

doc.add_paragraph(
    '基于 ColabFold (AlphaFold2) 的序列→结构预测。支持从序列生成高精度三维结构模型，'
    '并自动解析置信度指标。'
)

doc.add_heading('6.1 预测参数', level=2)
add_table_simple(doc,
    ['参数', '类型', '默认值', '说明'],
    [
        ['模型数', '滑块 (1-5)', '1', 'ColabFold 模型数量'],
        ['回收步数', '滑块 (1-10)', '1', '回收/精修步数'],
        ['使用 AMBER 精修', '复选框', 'False', '启用 Amber 能量最小化'],
        ['stop-at-score', '自动', '85', 'pLDDT 超过此值时提前终止'],
        ['模型类型', '下拉', 'auto', 'auto / alphafold2_ptm / alphafold2_multimer_v3'],
    ]
)

doc.add_heading('6.2 结果解读', level=2)
doc.add_paragraph(
    '预测完成后显示：\n'
    '• pLDDT — 逐残基置信度 (0-100)，>90 高置信、70-90 中等、<70 低置信\n'
    '• pTM — 全局折叠置信度 (0-1)，>0.8 可靠\n'
    '• ipTM — 界面置信度 (0-1)，多链结构时可用\n'
    '• max PAE — 最大预测对齐误差，越低越好\n'
    '• 预测的 PDB 文件可一键发送到设计流程'
)

add_note(doc, '多链结构 (抗体-抗原复合物) 应使用 alphafold2_multimer_v3 模型类型以获得 ipTM 评分。', 'tip')

doc.add_page_break()

# ═══════════════ 7. Settings ═══════════════
doc.add_heading('7. 系统设置（Tab 4 — ⚙️ 设置）', level=1)

doc.add_paragraph(
    '修改运行时参数，修改后即时生效。所有设置可在 app_config.yaml 中持久化。'
)
add_table_simple(doc,
    ['设置项', '说明', '默认值'],
    [
        ['AF2 可执行文件', 'colabfold_batch 路径', 'venv/Scripts/colabfold_batch'],
        ['AF2 模型数', '默认运行模型数', '1'],
        ['AF2 回收步数', '默认回收步数', '1'],
        ['AF2 stop-at-score', '提前终止分数阈值', '85'],
        ['BFN checkpoint', '模型权重路径', '(用户缓存目录)'],
        ['BFN config', '模型配置文件', 'configs/demo_design.yml'],
        ['服务器端口', 'Gradio 服务端口', '7860'],
        ['自动打开浏览器', '启动时自动打开', 'True'],
    ]
)

doc.add_page_break()

# ═══════════════ 8. Target Design ═══════════════
doc.add_heading('8. 靶点设计（Tab 5 — 🎯 靶点设计）', level=1)

doc.add_paragraph(
    '面向靶点蛋白的抗体/蛋白设计工作流。五步流程：靶点输入 → 表位分析 → 抗体设置 → '
    '约束设计 → 结果验证。'
)

doc.add_heading('8.1 步骤 1：靶点输入', level=2)
doc.add_paragraph(
    '两种方式提供靶点结构：\n'
    '• 直接上传靶点 PDB 文件\n'
    '• 输入 FASTA 序列 → AlphaFold2 预测结构\n'
    '选择需要分析的链 (Dropdown 自动检测可用链)'
)

doc.add_heading('8.2 步骤 2：表位分析', level=2)
doc.add_paragraph(
    '基于三种评分的复合表位预测：\n'
    '• SASA (溶剂可及性) — 暴露于溶剂的残基比例\n'
    '• 亲水性 (Hydrophilicity) — 亲水残基倾向\n'
    '• 凸出指数 (Protrusion) — 结构表面的突出程度\n'
    '\n'
    '参数：\n'
    '• Top-N：显示前 N 个候选表位残基 (5-50，默认 20)\n'
    '• 确认表位区域：手工调整后点击确认'
)

doc.add_heading('8.3 步骤 3：抗体设置', level=2)
doc.add_paragraph(
    '提供抗体模板：\n'
    '• 上传抗体 PDB (含 H/L 链)\n'
    '• 或通过 FASTA → PDB 生成 α-螺旋模板\n'
    '设置重链/轻链 ID (默认 H/L)'
)

doc.add_heading('8.4 步骤 4：约束设计', level=2)
doc.add_paragraph(
    '核心设计参数：\n'
    '• 设计工具：BFN (抗体CDR) / ProteinMPNN / ESM-IF\n'
    '• 约束模式：启用后仅设计与表位接触的残基 (默认开启)\n'
    '• 距离阈值：CA-CA 距离阈值 (Å)，确定面向表位的残基 (5-20 Å，默认 10)\n'
    '\n'
    '约束原理：计算抗体残基 CA 与表位残基 CA 的距离，仅对距离 < 阈值的残基进行设计。'
)

doc.add_heading('8.5 步骤 5：结果与验证', level=2)
doc.add_paragraph(
    '设计完成后显示：\n'
    '• 设计结果：序列、PPL、pLDDT、ipTM\n'
    '• 验证报告：接触分析、极性/疏水界面比例、H-bond 网络评估\n'
    '• 操作按钮：下载 FASTA、发送到 AF2 Multimer 预测、发送到单步设计'
)

doc.add_page_break()

# ═══════════════ 9. Iterative Refinement ═══════════════
doc.add_heading('9. 迭代精修（Tab 6 — 🔄 迭代精修）', level=1)

doc.add_paragraph(
    '借鉴 BindCraft 策略，通过 AF2 外部验证驱动多轮设计→验证→筛选→再设计循环，'
    '逐步提升设计的结构质量和置信度。'
)

doc.add_heading('9.1 工作原理', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '第 1 轮：设计 N 条序列 → AF2 验证每条 → 级联筛选 Top-K → 保留 AF2 结构\n'
    '第 2 轮：用 Top-K 的 AF2 结构作模板 → 再设计 → AF2 验证 → 筛选 → 更新 Top-K\n'
    '...重复至收敛 (ipTM 提升 < Δ) 或达到最大轮次'
)
run.font.size = Pt(10)

doc.add_heading('9.2 参数说明', level=2)
add_table_simple(doc,
    ['参数', '范围', '默认值', '说明'],
    [
        ['PDB 模板', '文件', '(必选)', '起始结构 PDB 文件'],
        ['设计区域', '文本', 'A:10-30', '链:起始-结束 格式'],
        ['每轮样本数', '3-20', '5', '每轮 BFN 生成的序列数量'],
        ['保留 Top-K', '1-5', '2', '每轮保留进入下一轮的最优序列数'],
        ['最大轮次', '1-5', '2', '最多迭代轮次'],
        ['AF2 回收步数', '0-3', '1', 'ColabFold 回收步数 (0=跳过 AF2)'],
    ]
)

add_note(doc, 'AF2 验证每条序列约需 5-15 分钟。5 条 × 2 轮 = 约 50-150 分钟。'
          '建议先用小样本数 (3) + 2 轮验证流程。', 'warn')

doc.add_heading('9.3 收敛条件', level=2)
doc.add_paragraph(
    '• 当前轮最优 ipTM 相比上轮提升 < 0.02 (默认阈值)\n'
    '• 或达到最大轮次限制\n'
    '• 或无可设计序列 (全部被级联过滤淘汰)'
)

doc.add_heading('9.4 输出', level=2)
doc.add_paragraph(
    '• 运行日志：每轮详细的 AF2 进度和中间结果\n'
    '• 精修报告：最终汇总 (最优 ipTM、pTM、序列、收敛状态)'
)

doc.add_page_break()

# ═══════════════ 10. Confidence Metrics ═══════════════
doc.add_heading('10. 置信度指标说明', level=1)

doc.add_heading('10.1 BFN 内置置信度', level=2)
doc.add_paragraph(
    'BFN 模型在采样过程中输出三个置信度指标 (通过 receiver.py 的置信度头):\n'
    '• pLDDT (predictive LDDT): 预测的逐残基 LDDT 值，通过 sigmoid 映射到 [0,1]\n'
    '• ipTM (interface predicted TM-score): 预测的界面 TM-score，通过 sigmoid 映射到 [0,1]\n'
    '• PAE (predicted aligned error): 预测的对齐误差矩阵'
)
add_note(doc, '当前 BFN 模型使用抗体 CDR 数据训练，置信度头对通用蛋白质输出接近 0。'
          '正在进行通用蛋白质微调以改善此问题。', 'warn')

doc.add_heading('10.2 AF2 置信度 (ColabFold)', level=2)
doc.add_paragraph(
    'ColabFold 运行后从 *_scores_rank_*.json 文件提取：\n'
    '• pLDDT: 0-100 标度，>90 为高置信\n'
    '• pTM: 0-1 标度，>0.8 为可靠折叠\n'
    '• ipTM: 0-1 标度，仅 multimer 模式计算\n'
    '• max_pae: 最大预测对齐误差 (Å)，<10 为良好'
)
add_note(doc, 'pTM 评估全局折叠质量，ipTM 评估多链界面质量。对于单链蛋白，关注 pTM + pLDDT。', 'tip')

doc.add_page_break()

# ═══════════════ 11. Cascade Filter ═══════════════
doc.add_heading('11. 级联过滤机制', level=1)

doc.add_paragraph(
    '级联过滤器对设计结果进行三阶段筛选和排序，确保输出最高质量的候选序列。'
)

doc.add_heading('11.1 BFN 级联过滤 (apply_cascade)', level=2)
doc.add_paragraph('三级处理流程：')
add_table_simple(doc,
    ['阶段', '操作', '默认阈值'],
    [
        ['第 1 级: 硬阈值', '筛选 pLDDT / ipTM / PPL / MPNN score', 'pLDDT ≥ -0.01, ipTM ≥ -0.01, PPL ≤ 100'],
        ['第 2 级: 去重', '相同序列仅保留 PPL 最小的', '—'],
        ['第 3 级: 综合评分', '加权组合: ipTM(0.35) + pLDDT(0.25) + PPL⁻¹(0.25) + recovery(0.15)', '—'],
    ]
)

doc.add_heading('11.2 AF2 级联过滤 (apply_cascade_af2)', level=2)
add_table_simple(doc,
    ['阶段', '操作', '默认阈值'],
    [
        ['第 1 级: AF2 阈值', '筛选 pLDDT / pTM / ipTM / max PAE', 'pLDDT ≥ 60, pTM ≥ 0.4, ipTM ≥ 0.3, max PAE ≤ 15'],
        ['第 2 级: 去重', '相同序列仅保留 pTM 最高的', '—'],
        ['第 3 级: 综合评分', '加权组合: ipTM(0.35) + pTM(0.25) + pLDDT(0.25) + PPL⁻¹(0.15)', '—'],
    ]
)

doc.add_heading('11.3 自定义阈值', level=2)
doc.add_paragraph(
    '在代码中调用时传入自定义阈值和权重：'
)
add_code_block(doc, 'from cascade_filter import apply_cascade_af2')
add_code_block(doc, 'filtered, report = apply_cascade_af2(')
add_code_block(doc, '    results,')
add_code_block(doc, '    thresholds={"plddt_min": 70, "ptm_min": 0.5},')
add_code_block(doc, '    weights={"iptm": 0.4, "ptm": 0.3, "plddt": 0.2, "ppl_inv": 0.1}')
add_code_block(doc, ')')

doc.add_page_break()

# ═══════════════ 12. Config Reference ═══════════════
doc.add_heading('12. 配置文件参考', level=1)

doc.add_paragraph('项目根目录下的 app_config.yaml 控制所有默认参数。主要配置段：')

doc.add_heading('12.1 [alphafold] — AlphaFold2 配置', level=2)
add_table_simple(doc,
    ['键', '类型', '说明'],
    [
        ['af2.venv', '路径', '虚拟环境/colabfold 根目录'],
        ['defaults.model_type', '字符串', 'auto / alphafold2_ptm / alphafold2_multimer_v3'],
        ['defaults.rank', '字符串', 'auto / plddt / ptm / iptm / multimer'],
        ['defaults.num_recycle', '整数', '默认回收步数'],
        ['output_dir', '路径', 'AF2结果输出目录'],
    ]
)

doc.add_heading('12.2 [models] — 模型路径', level=2)
add_table_simple(doc,
    ['键', '说明'],
    [
        ['models.bfn.checkpoint', 'BFN 模型权重 .pt 文件路径'],
        ['models.bfn.config', 'BFN 配置文件 yml 路径'],
        ['models.esmif.model_name', 'ESM-IF 模型名称 (HuggingFace)'],
        ['models.proteinmpnn.*', 'ProteinMPNN 权重目录'],
    ]
)

doc.add_heading('12.3 [bfn_defaults] — BFN 默认参数', level=2)
add_table_simple(doc,
    ['键', '说明'],
    [
        ['bfn_defaults.antibody.cdrs', '默认 CDR 环区列表'],
        ['bfn_defaults.antibody.num_samples', '抗体默认样本数'],
        ['bfn_defaults.protein.region', '通用蛋白默认设计区域'],
        ['bfn_defaults.protein.num_samples', '通用蛋白默认样本数'],
    ]
)

doc.add_heading('12.4 [target_design] — 靶点设计配置', level=2)
add_table_simple(doc,
    ['键', '说明'],
    [
        ['target_design.epitope.*', '表位评分权重 (SASA/亲水性/凸出指数)'],
        ['target_design.design.*', '约束设计默认参数'],
        ['target_design.contact.*', '接触分析阈值'],
    ]
)

doc.add_page_break()

# ═══════════════ 13. CLI Tools ═══════════════
doc.add_heading('13. 命令行工具', level=1)

doc.add_heading('13.1 命令行设计', level=2)
doc.add_paragraph('通过命令行直接运行设计，无需启动 Gradio 界面：')
add_code_block(doc, 'python design_protein.py --pdb input.pdb --region A:10-30')
add_code_block(doc, 'python design_seq.py --fasta sequence.fasta')

doc.add_heading('13.2 验证脚本', level=2)
doc.add_paragraph('验证整个管线是否正常工作：')
add_code_block(doc, 'python validate_8proteins_v2.py')
doc.add_paragraph('此脚本运行 7 个阶段的自动化测试，覆盖模块导入、级联过滤、BFN 设计、'
                   'AF2 验证、迭代精修和 UI 集成检查。')

doc.add_heading('13.3 AF2 验证器', level=2)
doc.add_paragraph('独立使用 AF2 批量验证器：')
add_code_block(doc, 'from af2_validator import validate_sequences')
add_code_block(doc, 'results = validate_sequences(["SEQUENCE1", "SEQUENCE2"])')
add_code_block(doc, '# 返回列表: [{sequence, plddt, ptm, iptm, max_pae, pdb_path}, ...]')

doc.add_heading('13.4 迭代精修 API', level=2)
add_code_block(doc, 'from iterative_refiner import IterativeRefiner')
add_code_block(doc, 'refiner = IterativeRefiner(')
add_code_block(doc, '    design_fn=my_design_function,')
add_code_block(doc, '    af2_validator=validate_sequences,')
add_code_block(doc, '    cascade_fn=apply_cascade_af2,')
add_code_block(doc, '    n_samples=10, top_k=3, max_rounds=3')
add_code_block(doc, ')')
add_code_block(doc, 'result = refiner.run("input.pdb", "A:10-30")')

doc.add_page_break()

# ═══════════════ 14. Reference Results ═══════════════
doc.add_heading('14. 8种测试蛋白质设计结果', level=1)

doc.add_paragraph(
    '以下为 8 种代表性蛋白质使用 BFN (通用蛋白质模式, 5 样本, 3 次回收) 的设计结果。'
    '设计区域为序列中间约 50 个残基。'
)

doc.add_heading('14.1 蛋白质序列', level=2)
proteins = {
    'FABP3': 'MVDAFLGTWKLVDSKNFDDYMKSLGVGFATRQVASMTKPTTIIEKNGDILTLKTHSTFKNTEISFKLGVEFDETTADDRKVKSIVTLDGGKLVHLQKWDGQET',
    'ENO2': 'MSIEKIWAREILDSRGNPTVEVDLYTAKGLFRAAVPSGASTGIYEALELRDGDKGRYLGKGVLKAVENINNTLGPALLQKKLSVVDQEKVDKFMIELDG',
    'ENO1': 'MSILKIHAREIFDSRGNPTVEVDLFTSKGLFRAAVPSGASTGIYEALELRDNDKTRYMGKGVSRAVEHINKTIAPALVSKKLNVTEQEKIDKLMIEMDG',
    'MAPT': 'MAEPRQEFEVMEDHAGTYGLGDRKDQGGYTMHQDQEGDTDAGLKESPLQTPTEDGSEEPGSETSDAKSTPTAEDVTAPLVDEGAPGKQAAAQPHTEIPEG',
    'NRGN': 'MDCCTENACSKPDDDILDIPLDDPGANAAAAKIQASFRGHMARKKIKSGECGRKGPGPGGPGGAGGARGGAGGGPSGD',
    'MIF': 'MPMFIVNTNVPRASVPDGFLSELTQQLAQATGKPPQYIAVHVVPDQLMAFGGSSEPCALCSLHSIGKIGGAQNRSYSKLLCGLLAERLRISPDRVYINYY',
    'TMSB10': 'MADKPDMGEIASFDKAKLKKTETQEKNTLPTKETIEQEKRSEIS',
    'GLOD4': 'MAAVQALEVLKEQGLVQLRAQGTGTSNGLTAQKHYLLGNVLKPNKGSGVQGWRVGSVFHQDPENPSLFLGQGGQCVSLWGRDVPGSPAAGALQAPGDL',
}
for name, seq in proteins.items():
    doc.add_paragraph(f'{name} ({len(seq)} aa):', style='List Bullet')
    add_code_block(doc, seq)

doc.add_heading('14.2 设计结果摘要', level=2)
doc.add_paragraph('使用 BFN 通用蛋白质模式，每蛋白 5 样本，3 次回收，不同随机种子。')

add_table_simple(doc,
    ['蛋白质', '长度', '设计区域', '最佳PPL', '平均PPL', 'BFN pLDDT', 'BFN ipTM'],
    [
        ['FABP3', '103', '50', '37.25', '42.16', '0.000', '0.000'],
        ['ENO2', '99', '50', '29.19', '39.09', '0.000', '0.000'],
        ['ENO1', '99', '50', '32.95', '48.24', '0.000', '0.000'],
        ['MAPT', '100', '50', '29.43', '39.36', '0.000', '0.000'],
        ['NRGN', '78', '50', '30.47', '46.34', '0.000', '0.000'],
        ['MIF', '100', '50', '40.65', '53.24', '0.000', '0.000'],
        ['TMSB10', '44', '44', '26.39', '36.17', '0.000', '0.000'],
        ['GLOD4', '98', '50', '40.77', '45.84', '0.000', '0.000'],
    ]
)

add_note(doc, 'BFN pLDDT/ipTM = 0.000 是因为置信度头仅在抗体CDR数据上训练，对通用蛋白质未校准。'
          '建议以 PPL 作为主要筛选指标，并使用 AF2 进行独立验证。', 'warn')

doc.add_heading('14.3 输出文件说明', level=2)
doc.add_paragraph(
    '所有结果文件位于 Desktop/protein/ 目录：\n'
    '• original_8proteins.fasta — 8种蛋白质的原始序列\n'
    '• summary.tsv — 批次汇总表\n'
    '• all_results.json — 完整设计结果 (含每条序列的详细指标)\n'
    '• [蛋白质名]_designs.fasta — FASTA 格式的设计序列 (含排名和置信度标注)\n'
    '• [蛋白质名]_results.json — 各蛋白质的结构化详细结果'
)

doc.add_heading('14.4 各蛋白质最优设计序列', level=2)
best_seqs = {
    'FABP3': 'WIAATKAYRIWKIPLFVVGENFKTNGVVFSAPQSRCQPSAVQFAKRNHDR',
    'ENO2': 'TDYRTTMRQWGSAPVNALGRDYNNVAVHDEGHPKLLHSGMHHSSYFSLAF',
    'ENO1': 'HMNWDSTPNVGDVVWPPAVWLPKPEVAVIYWPRSRCNSGSCAKHNYAKVQ',
    'MAPT': 'FFCTKVYHTVVGYFMRHAIYFAIPALVYSIPYVHPCVCVYYIVCHVCSIC',
    'NRGN': 'PQTFGMTQGQDYERLDRDVHHSELTDMEAKVCFFDVFGSHRLGNVDTSVD',
    'MIF': 'CTVYCHHDVFQNGNIQCLDWNLRMCQTWFQGGGFYVVYDCKNTKGGHGSV',
    'TMSB10': 'GWMLTMLKTFYDYVPDDNLSGCVFARKWIVKFDTTCGYYDKKSP',
    'GLOD4': 'RDAAYGPAEYQLGEAYSNGVHEFSCVVLAGMSRPKKNEYKVVKQGQVWNH',
}
for name, seq in best_seqs.items():
    doc.add_paragraph(f'{name} (最佳PPL):', style='List Bullet')
    add_code_block(doc, seq)

doc.add_heading('14.5 设计区域说明', level=2)
doc.add_paragraph(
    '对于长度 > 50 aa 的蛋白质 (FABP3, ENO2, ENO1, MAPT, MIF, GLOD4)，'
    '取中间约 50 个残基作为设计区域。对于较短的蛋白质 (NRGN 78 aa, TMSB10 44 aa)，'
    '使用全长序列进行设计。所有模板均为标准 α-螺旋骨架几何。'
)
doc.add_paragraph(
    '注意: 由于所有 50-aa 模板具有相同的 α-螺旋骨架结构 (仅序列不同)，'
    'BFN 的设计输出不依赖于输入序列的具体氨基酸，而取决于骨架几何和随机种子。'
    '输入序列仅用于确定设计区域的长度。'
)

# ═══════════════ Footer ═══════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('蛋白质序列设计平台 · BFN + ProteinMPNN + ESM-IF + AlphaFold2\n'
                 '如有问题请联系开发团队')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ──
doc.save(OUT_PATH)
print(f'Word 文档已保存至: {OUT_PATH}')
